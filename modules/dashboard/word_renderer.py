from docx import Document
import os

def generate_oficio_word(out_dir: str, ciudad: str, fecha: str, asunto: str, referencia: str, destinatario: str, cargo_dest: str, entidad_dest: str, numero_oficio: str, resultado: str, institucion_nombre: str, firmante_nombre: str, firmante_cargo: str, usuario_actual: str, rol_actual: str, ciudadano_line: str, year_fiscal: str = "AÑO DE LA RECUPERACIÓN Y CONSOLIDACIÓN DE LA ECONOMÍA PERUANA") -> str:
    doc = Document()
    doc.add_heading(f"Oficio N° {numero_oficio}", 0)
    doc.add_paragraph(f"{ciudad}, {fecha}")
    doc.add_paragraph(f"Año Fiscal: {year_fiscal}")
    doc.add_paragraph(f"Institución: {institucion_nombre}")
    doc.add_paragraph(f"Señor(a): {destinatario} {cargo_dest} {entidad_dest}")
    doc.add_paragraph(f"Asunto: {asunto}")
    doc.add_paragraph(f"Referencia: {referencia}")
    doc.add_paragraph("")
    doc.add_paragraph("Tengo el agrado de dirigirme a usted para expresarle un cordial saludo y, en atención al documento de la referencia, informarle lo siguiente:")
    doc.add_paragraph("En la verificación realizada en nuestros registros sobre la existencia de ficha de inscripción militar de los ciudadanos señalados en el documento de la referencia, se ha procedido a realizar la búsqueda, con el resultado siguiente:")
    doc.add_paragraph(f"{ciudadano_line}")
    if resultado.upper() == "POSITIVO":
        doc.add_paragraph("Por lo tanto, se deja constancia de la inscripción militar del ciudadano indicado dentro de los registros de esta Oficina de Registro Militar Departamental N.° 055-A, adjuntándose al presente oficio la documentación sustentatoria correspondiente.")
    else:
        doc.add_paragraph("Por lo tanto, no existe constancia de inscripción militar a nombre del ciudadano indicado dentro de los registros físicos ni digitales de esta Oficina de Registro Militar Departamental N.° 055-A.")
    doc.add_paragraph("")
    doc.add_paragraph("Hago propicia la oportunidad para expresarle las seguridades de mi especial consideración y estima personal.")
    doc.add_paragraph("Dios guarde a Ud.")
    doc.add_paragraph(f"{firmante_nombre}")
    doc.add_paragraph(f"{firmante_cargo}")
    # Limpiar nombre de archivo para evitar caracteres no válidos
    safe_numero_oficio = str(numero_oficio).replace("/", "-").replace("\\", "-").replace(":", "-").replace("*", "-").replace("?", "-").replace("\"", "-").replace("<", "-").replace(">", "-").replace("|", "-")
    out_path = os.path.join(out_dir, f"oficio_{safe_numero_oficio}.docx")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path
